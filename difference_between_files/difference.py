import difflib
import io
import re
import logging
from docx import Document
from docx.enum.text import WD_COLOR_INDEX
from docx.shared import Inches
from string import punctuation

from difference_between_files.acceptable import replacements, skips

logging.basicConfig(level=logging.DEBUG)
logger = logging.getLogger('Documents')


class DiffData:
    def __init__(self, current_number, last_known_number, first_column, second_column, is_different):
        self.current_number = current_number
        self.last_known_number = last_known_number
        self.first_column = first_column
        self.second_column = second_column
        self.is_different = is_different


def list_from_string(document: str) -> list:
    lst = [i.replace('®', '').replace('', '') for i in document.split('\n\n')]
    return lst


def get_diff(list1: list, list2: list) -> list:
    last_known_number = ''
    for text1, text2 in zip(list1, list2):
        diff_items = difflib.ndiff([text1], [text2])
        logger.info('Find error')
        for diff in diff_items:
            diff = diff.strip('®')
            is_different = diff[0] in ('+', '-')
            break
        match_number = re.search(r"^(\.?,?\d{0,2}){0,4} ?", text1)
        current_number = match_number[0] if match_number and not match_number[0].isspace() else ""
        last_known_number = current_number if current_number else last_known_number
        diffs_data = DiffData(current_number=current_number, last_known_number=last_known_number, first_column=text1,
                              second_column=text2, is_different=is_different)

        yield diffs_data


def filter_diffs(diffs):
    duplicate = []
    for num, diff in enumerate(diffs):
        if not diff[1]:
            result = diff[0][:7]
            for n, dif in enumerate(diffs[num:]):
                number = n + num
                if result == dif[1][:7]:
                    duplicate.append((num, number, (diffs[number][0], diffs[num][1])))
                    break

    duplicate.reverse()

    for i in duplicate:
        diffs.pop(i[1])
        diffs[i[0]] = i[2]
    return diffs


def save_disagreement(file1: str, file2: str, count_error: int) -> io.BytesIO:
    logger.info('Create document')
    result = Document()
    result.add_heading("Протокол разногласий")
    table = result.add_table(rows=1, cols=3)
    table.style = "TableGrid"
    table.autofit = False
    heading_cells = table.rows[0].cells
    heading_cells[0].text = "№"
    heading_cells[0].width = Inches(0.6)
    heading_cells[1].text = "Редакция заказчика"
    heading_cells[1].width = Inches(3)
    heading_cells[2].text = "Редакция исполнителя"
    heading_cells[2].width = Inches(3)

    list1, list2 = list_from_string(file1), list_from_string(file2)
    diffs = list(get_diff(list1, list2))
    number_flag = ''
    for diff in diffs:
        if not diff.last_known_number and not diff.current_number:
            number = ''
        else:
            number = diff.last_known_number + ' ✓ ' if not diff.current_number else diff.current_number
        text1, text2 = [re.sub(r"(?:^(\.?,?\d{0,2}){0,4} ?|\.?,?$)", "", text).strip() for text in
                        [diff.first_column, diff.second_column]]
        if not diff.is_different and count_error == 0:
            cells = table.add_row().cells
            cells[0].width = Inches(0.6)
            cells[1].width = Inches(3)
            cells[2].width = Inches(3)
            cells[0].text = number
            paragraph1 = cells[1].paragraphs[0]
            paragraph2 = cells[2].paragraphs[0]
            paragraph1.add_run(text1)
            paragraph2.add_run(text2)
            continue
        if not diff.is_different:
            continue
        cells = table.add_row().cells
        cells[0].width = Inches(0.6)
        cells[1].width = Inches(3)
        cells[2].width = Inches(3)
        cells[0].text = number

        left_diff_count = 0
        right_diff_count = 0
        sequence = difflib.SequenceMatcher(a=text1.lower(), b=text2.lower(), autojunk=False)

        paragraph1 = cells[1].paragraphs[0]
        paragraph2 = cells[2].paragraphs[0]

        logger.info('Write document')
        for op, i1, i2, j1, j2 in sequence.get_opcodes():
            run1 = paragraph1.add_run(text1[i1:i2])
            run2 = paragraph2.add_run(text2[j1:j2])
            if op in ("delete", "insert"):
                if text1[i1:i2] not in skips:
                    run1.font.highlight_color = WD_COLOR_INDEX.YELLOW
                    left_diff_count += i2 - i1
                if text2[j1:j2] not in skips:
                    run2.font.highlight_color = WD_COLOR_INDEX.YELLOW
                    right_diff_count += j2 - j1
            if op == "replace":
                if (text1[i1:i2], text2[j1:j2]) not in replacements and (
                        text2[j1:j2],
                        text1[i1:i2],
                ) not in replacements:
                    run1.font.highlight_color = WD_COLOR_INDEX.YELLOW
                    run2.font.highlight_color = WD_COLOR_INDEX.YELLOW
                    left_diff_count += i2 - i1
                    right_diff_count += j2 - j1

        if left_diff_count <= count_error and right_diff_count <= count_error:
            table._tbl.remove(table.rows[-1]._tr)

    file_stream = io.BytesIO()
    result.save(file_stream)
    file_stream.seek(0)

    return file_stream
